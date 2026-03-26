"""
export_journal_docx.py
Generate FIELD_JOURNAL.docx and FIELD_JOURNAL_DEEP_ONLY.docx from
backend/data/journal/journal.json.

Usage:
    py scripts/export_journal_docx.py
    py scripts/export_journal_docx.py --deep-only
    py scripts/export_journal_docx.py --output-dir docs/

Outputs:
    docs/FIELD_JOURNAL.docx
    docs/FIELD_JOURNAL_DEEP_ONLY.docx
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime, timezone
from collections import defaultdict

sys.path.insert(0, "C:/Users/SeanP/OneDrive - mapleclose.ca/Sean/AI/LLEv2/Lib/site-packages")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

JOURNAL_PATH = "backend/data/journal/journal.json"
OUTPUT_DIR   = "docs"

DEEP_THRESHOLD    = 0.70
SHALLOW_THRESHOLD = 0.50

ROGET_CLASS_NAMES = {
    "1": "Abstract Relations",
    "2": "Space",
    "3": "Matter",
    "4": "Intellect",
    "5": "Volition",
    "6": "Affections",
}

# Entry title colours by class (higher-numbered class wins)
ROGET_CLASS_COLOURS = {
    "1": RGBColor(0x00, 0xb4, 0xd8),   # cyan
    "2": RGBColor(0xe0, 0x40, 0xa0),   # magenta
    "3": RGBColor(0xf0, 0x70, 0x20),   # orange
    "4": RGBColor(0x4e, 0xcb, 0x71),   # green
    "5": RGBColor(0xa0, 0x70, 0xe0),   # violet
    "6": RGBColor(0xe0, 0x50, 0x50),   # warm red
}

GREY_LIGHT  = RGBColor(0x88, 0x88, 0x88)
GREY_MED    = RGBColor(0x44, 0x44, 0x44)
GREY_DARK   = RGBColor(0x33, 0x33, 0x33)
GREY_SUB    = RGBColor(0x66, 0x66, 0x66)
GOLD        = RGBColor(0xff, 0xd7, 0x00)
BLACK       = RGBColor(0x00, 0x00, 0x00)

AUTHOR = "Claude (Anthropic) · Sean Patrick Morris"
SITE   = "datasculptures.com"

# ---------------------------------------------------------------------------
# Security: keep output inside project
# ---------------------------------------------------------------------------

def safe_output_path(filename: str) -> str:
    project_root = os.path.abspath(os.path.dirname(os.path.dirname(__file__)))
    out = os.path.abspath(os.path.join(project_root, OUTPUT_DIR, filename))
    if not out.startswith(project_root):
        raise ValueError(f"Output path escapes project directory: {out}")
    return out

# ---------------------------------------------------------------------------
# Low-level docx helpers (from generate_docs_docx.py pattern)
# ---------------------------------------------------------------------------

def _add_page_numbers(section):
    footer = section.footer
    para   = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    run = para.add_run(f"{SITE}   ·   Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = GREY_LIGHT

    for field_text in ("PAGE", "NUMPAGES"):
        if field_text == "NUMPAGES":
            r = para.add_run(" of ")
            r.font.size = Pt(9)
            r.font.color.rgb = GREY_LIGHT
        run_f = para.add_run()
        run_f.font.size = Pt(9)
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run_f._r.append(fld)
        instr = OxmlElement("w:instrText")
        instr.text = field_text
        run_f._r.append(instr)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        run_f._r.append(fld2)


def _set_header(section):
    header = section.header
    para   = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run(AUTHOR)
    run.font.size = Pt(9)
    run.font.color.rgb = GREY_LIGHT


def _set_margins(section):
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)


def _para_spacing(para, before=0, after=4):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"),  str(after))
    pPr.append(sp)


def _add_thin_rule(doc):
    """Add a 0.5pt #cccccc horizontal rule paragraph."""
    para = doc.add_paragraph()
    _para_spacing(para, before=4, after=4)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")      # 0.5pt
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_simple_table(doc, headers, rows, col_widths=None):
    """Add a plain table with a shaded header row."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    # Header
    hrow = table.rows[0]
    for j, h in enumerate(headers):
        cell = hrow.cells[j]
        cell.text = ""
        run  = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        # Light shade
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "DDEEFF")
        tcPr.append(shd)
    # Data rows
    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ""
            run  = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Entry parsing helpers
# ---------------------------------------------------------------------------

def _parse_pair(entry: dict) -> tuple[str, str]:
    """Return (term_a, term_b) from user_notes 'X vs Y', else (user_notes, '')."""
    notes = (entry.get("user_notes") or "").strip()
    m = re.match(r"^(.+?)\s+vs\s+(.+)$", notes, re.IGNORECASE)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return notes, ""


def _parse_classes(entry: dict) -> tuple[str, str]:
    """Return (class_id_a, class_id_b) from tags, lower first."""
    tags = entry.get("tags") or []
    class_ids = sorted(
        t.replace("class_", "") for t in tags if re.match(r"^class_[1-6]$", t)
    )
    if len(class_ids) >= 2:
        return class_ids[0], class_ids[1]
    if len(class_ids) == 1:
        return class_ids[0], ""
    # Fall back to roget_context
    ctx = entry.get("roget_context") or {}
    ca  = str(ctx.get("class_a", "")).strip()
    cb  = str(ctx.get("class_b", "")).strip()
    ids = sorted(c for c in [ca, cb] if c and c in ROGET_CLASS_NAMES)
    if len(ids) >= 2:
        return ids[0], ids[1]
    if len(ids) == 1:
        return ids[0], ""
    return "", ""


def _title_colour(class_a: str, class_b: str) -> RGBColor:
    """Return colour for the higher-numbered (deeper) Roget class."""
    for cls in reversed(["1","2","3","4","5","6"]):
        if cls == class_b or cls == class_a:
            return ROGET_CLASS_COLOURS.get(cls, BLACK)
    return BLACK


def _level_tag(entry: dict) -> str:
    tags = entry.get("tags") or []
    for lev in ("cross_class", "cross_section", "adjacent_cat"):
        if lev in tags:
            return lev
    return ""


def _fmt_date(ts: str) -> str:
    try:
        dt = datetime.fromisoformat(ts)
        return dt.strftime("%Y-%m-%d")
    except Exception:
        return ts[:10] if ts else ""


# ---------------------------------------------------------------------------
# Entry rendering
# ---------------------------------------------------------------------------

def render_entry(doc: Document, entry: dict, first: bool = False) -> None:
    if not first:
        _add_thin_rule(doc)

    desert = entry.get("desert_value") or 0.0
    term_a, term_b = _parse_pair(entry)
    class_a, class_b = _parse_classes(entry)
    colour = _title_colour(class_a, class_b)

    # ── Title line: "Term A ↔ Term B   [desert]" ──────────────────────
    title_para = doc.add_paragraph()
    _para_spacing(title_para, before=6, after=1)

    title_text = f"{term_a} ↔ {term_b}" if term_b else term_a
    run_title  = title_para.add_run(title_text)
    run_title.bold       = True
    run_title.font.size  = Pt(11)
    run_title.font.color.rgb = colour

    # Right-align the desert value using a tab stop trick via spaces
    run_tab = title_para.add_run("   ")
    run_tab.font.size = Pt(11)

    run_dv = title_para.add_run(f"{desert:.4f}")
    run_dv.bold       = True
    run_dv.font.size  = Pt(11)
    run_dv.font.color.rgb = colour

    # ── Roget class line ───────────────────────────────────────────────
    name_a = ROGET_CLASS_NAMES.get(class_a, "")
    name_b = ROGET_CLASS_NAMES.get(class_b, "")
    date_s = _fmt_date(entry.get("timestamp", ""))

    if name_a and name_b:
        class_text = f"{name_a} ↔ {name_b}"
    elif name_a:
        class_text = name_a
    else:
        class_text = ""

    if class_text or date_s:
        sub_para = doc.add_paragraph()
        _para_spacing(sub_para, before=0, after=2)
        if class_text:
            r = sub_para.add_run(class_text)
            r.italic          = True
            r.font.size       = Pt(9)
            r.font.color.rgb  = GREY_SUB
        if date_s:
            r2 = sub_para.add_run(f"   {date_s}")
            r2.italic         = True
            r2.font.size      = Pt(9)
            r2.font.color.rgb = GREY_LIGHT

    # ── Nearest concepts ───────────────────────────────────────────────
    nc = entry.get("nearest_concepts") or []
    labels = ["Deepest near:", "2nd nearest: ", "3rd nearest: "]
    for i, concept in enumerate(nc[:3]):
        term = concept.get("term", "")
        dist = concept.get("distance")
        dist_s = f"{dist:.4f}" if isinstance(dist, float) else "—"
        line = f"{labels[i]}  {term}  (distance: {dist_s})"
        p = doc.add_paragraph()
        _para_spacing(p, before=0, after=1)
        r = p.add_run(line)
        r.font.name  = "Courier New"
        r.font.size  = Pt(9)
        r.font.color.rgb = GREY_MED

    # ── Generated description ──────────────────────────────────────────
    desc = (entry.get("generated_description") or "").strip()
    if desc:
        p = doc.add_paragraph()
        _para_spacing(p, before=2, after=2)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(desc)
        r.italic         = True
        r.font.size      = Pt(10)
        r.font.color.rgb = GREY_DARK

    # ── User notes (only if not the auto-generated "A vs B" note) ──────
    notes = (entry.get("user_notes") or "").strip()
    auto_note = f"{term_a} vs {term_b}".lower() if term_b else ""
    if notes and notes.lower() != auto_note:
        # Also skip if it's *only* the pair pattern with no extra text
        m = re.match(r"^(.+?)\s+vs\s+(.+)$", notes, re.IGNORECASE)
        if not m:
            p = doc.add_paragraph()
            _para_spacing(p, before=2, after=2)
            p.paragraph_format.left_indent = Cm(0.5)
            r = p.add_run(notes)
            r.font.size = Pt(10)

    # ── Tags / fabrication / starred ───────────────────────────────────
    tags = [t for t in (entry.get("tags") or []) if not re.match(r"^class_", t) and t not in ("cross_class","cross_section","adjacent_cat")]
    fab_status = (entry.get("fabrication_notes") or {}).get("status", "idea")
    starred    = entry.get("starred", False)

    meta_parts = []
    if tags:
        meta_parts.append("Tags: " + " · ".join(tags))
    if fab_status and fab_status != "idea":
        meta_parts.append(f"Fabrication: {fab_status}")

    if meta_parts or starred:
        p = doc.add_paragraph()
        _para_spacing(p, before=1, after=2)
        if meta_parts:
            r = p.add_run("  ".join(meta_parts))
            r.font.size      = Pt(8)
            r.font.color.rgb = GREY_LIGHT
        if starred:
            r2 = p.add_run("   ★ STARRED")
            r2.bold          = True
            r2.font.size     = Pt(9)
            r2.font.color.rgb = GOLD


# ---------------------------------------------------------------------------
# Section rendering
# ---------------------------------------------------------------------------

def render_section(doc: Document, title: str, subtitle: str, entries: list) -> None:
    h = doc.add_heading(title, level=1)
    _para_spacing(h, before=12, after=4)

    sub = doc.add_paragraph(subtitle)
    _para_spacing(sub, before=0, after=8)
    for r in sub.runs:
        r.italic         = True
        r.font.size      = Pt(10)
        r.font.color.rgb = GREY_SUB

    for i, entry in enumerate(entries):
        render_entry(doc, entry, first=(i == 0))


# ---------------------------------------------------------------------------
# Appendix A — statistics
# ---------------------------------------------------------------------------

def render_appendix_a(doc: Document, entries: list) -> None:
    doc.add_heading("Appendix A — Statistics", level=1)

    probe   = [e for e in entries if e.get("type") == "probe_discovery"]
    manual  = [e for e in entries if e.get("type") == "manual"]
    v1imp   = [e for e in entries if e.get("type") == "v1_import"]
    deep    = [e for e in probe if (e.get("desert_value") or 0) >= DEEP_THRESHOLD]
    shallow = [e for e in probe if SHALLOW_THRESHOLD <= (e.get("desert_value") or 0) < DEEP_THRESHOLD]

    # Summary paragraph
    p = doc.add_paragraph()
    lines = [
        f"Total entries: {len(entries)}",
        f"By type:  probe_discovery {len(probe)},  manual {len(manual)},  v1_import {len(v1imp)}",
        f"Deep (≥ {DEEP_THRESHOLD}): {len(deep)}    Shallow ({SHALLOW_THRESHOLD}–{DEEP_THRESHOLD}): {len(shallow)}",
    ]
    for ln in lines:
        r = p.add_run(ln + "\n")
        r.font.size = Pt(10)
    _para_spacing(p, before=4, after=8)

    # Level breakdown with mean desert
    levels_count: dict[str, int]   = defaultdict(int)
    levels_deserts: dict[str, list] = defaultdict(list)
    for e in probe:
        lv = _level_tag(e)
        if lv:
            levels_count[lv] += 1
            dv = e.get("desert_value")
            if isinstance(dv, float):
                levels_deserts[lv].append(dv)

    doc.add_heading("By level", level=2)
    lev_rows = []
    for k in sorted(levels_count.keys()):
        vals = levels_deserts[k]
        mean_s = f"{sum(vals)/len(vals):.4f}" if vals else "—"
        lev_rows.append((k, str(levels_count[k]), mean_s))
    _add_simple_table(doc, ["Level", "Count", "Mean desert"], lev_rows)

    # Class-pair table
    doc.add_heading("By Roget class pair (probe discoveries)", level=2)
    pair_data: dict[tuple, list] = defaultdict(list)
    for e in probe:
        ca, cb = _parse_classes(e)
        dv = e.get("desert_value")
        if isinstance(dv, float):
            key = (ROGET_CLASS_NAMES.get(ca, ca), ROGET_CLASS_NAMES.get(cb, cb))
            pair_data[key].append(dv)

    pair_rows = []
    for (na, nb), vals in pair_data.items():
        mean_d = sum(vals) / len(vals)
        pair_rows.append((na, nb, str(len(vals)), f"{mean_d:.4f}"))
    pair_rows.sort(key=lambda r: float(r[3]), reverse=True)
    _add_simple_table(doc, ["Class A", "Class B", "Count", "Mean desert"], pair_rows)

    # Top 10 deepest
    doc.add_heading("Top 10 deepest entries", level=2)
    top10 = sorted(
        [e for e in probe if isinstance(e.get("desert_value"), float)],
        key=lambda e: e["desert_value"],
        reverse=True,
    )[:10]
    top_rows = []
    for e in top10:
        ta, tb = _parse_pair(e)
        nc0  = (e.get("nearest_concepts") or [{}])[0]
        near = nc0.get("term", "—")
        desc = (e.get("generated_description") or "").strip()
        desc_preview = desc[:60] + ("…" if len(desc) > 60 else "") if desc else ""
        top_rows.append((ta, tb, f"{e['desert_value']:.4f}", near, desc_preview))
    _add_simple_table(doc, ["Term A", "Term B", "Desert", "Deepest near", "Description"], top_rows)


# ---------------------------------------------------------------------------
# Appendix B — fabrication queue
# ---------------------------------------------------------------------------

def render_appendix_b(doc: Document, entries: list) -> None:
    doc.add_heading("Appendix B — Fabrication Queue", level=1)
    fab_entries = [e for e in entries if (e.get("fabrication_notes") or {}).get("status", "idea") != "idea"]

    if not fab_entries:
        p = doc.add_paragraph("No entries marked for fabrication yet.")
        _para_spacing(p, before=4, after=4)
        for r in p.runs:
            r.italic         = True
            r.font.color.rgb = GREY_LIGHT
        return

    for i, entry in enumerate(fab_entries):
        render_entry(doc, entry, first=(i == 0))


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------

def render_title_page(doc: Document, entries: list) -> None:
    probe      = [e for e in entries if e.get("type") == "probe_discovery"]
    deep       = [e for e in probe if (e.get("desert_value") or 0) >= DEEP_THRESHOLD]
    shallow    = [e for e in probe if SHALLOW_THRESHOLD <= (e.get("desert_value") or 0) < DEEP_THRESHOLD]
    max_dv     = max((e.get("desert_value") or 0) for e in probe) if probe else 0
    n_desc     = sum(1 for e in entries if e.get("generated_description"))
    n_starred  = sum(1 for e in entries if e.get("starred"))

    for _ in range(4):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Field Journal")
    r.bold = True
    r.font.size = Pt(28)
    _para_spacing(t, before=0, after=4)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Latent Language Explorer V2")
    r2.bold = True
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0x00, 0xb4, 0xd8)
    _para_spacing(t2, before=0, after=8)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Discoveries from the Deserts of Vector Space")
    rs.italic        = True
    rs.font.size     = Pt(13)
    rs.font.color.rgb = GREY_SUB
    _para_spacing(sub, before=0, after=16)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = date_p.add_run("March 2026")
    rd.font.size     = Pt(11)
    rd.font.color.rgb = GREY_LIGHT
    _para_spacing(date_p, before=0, after=8)

    stat_p = doc.add_paragraph()
    stat_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stat_text = (
        f"{len(entries)} total entries\n"
        f"{len(deep)} deep (≥ {DEEP_THRESHOLD})  ·  "
        f"{len(shallow)} shallow ({SHALLOW_THRESHOLD}–{DEEP_THRESHOLD})\n"
        f"{n_desc} with descriptions  ·  {n_starred} starred\n"
        f"Max desert depth: {max_dv:.4f}"
    )
    rs2 = stat_p.add_run(stat_text)
    rs2.font.size     = Pt(10)
    rs2.font.color.rgb = GREY_MED
    _para_spacing(stat_p, before=0, after=0)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Main build function
# ---------------------------------------------------------------------------

def build_journal_doc(
    entries: list,
    output_path: str,
    deep_only: bool = False,
) -> int:
    """Build the .docx and return the number of entries written."""

    probe   = [e for e in entries if e.get("type") == "probe_discovery" and isinstance(e.get("desert_value"), float)]
    manual  = [e for e in entries if e.get("type") != "probe_discovery" or not isinstance(e.get("desert_value"), float)]
    # Include below-gate probe entries in manual section
    manual += [e for e in probe if e["desert_value"] < SHALLOW_THRESHOLD]

    deep_entries    = sorted([e for e in probe if e["desert_value"] >= DEEP_THRESHOLD],
                             key=lambda e: e["desert_value"], reverse=True)
    shallow_entries = sorted([e for e in probe if SHALLOW_THRESHOLD <= e["desert_value"] < DEEP_THRESHOLD],
                             key=lambda e: e["desert_value"], reverse=True)
    other_entries   = sorted(manual, key=lambda e: e.get("timestamp",""), reverse=True)

    doc = Document()
    section = doc.sections[0]
    _set_margins(section)
    _set_header(section)
    _add_page_numbers(section)

    render_title_page(doc, entries)

    render_section(
        doc,
        "Deep Discoveries",
        f"Desert depth ≥ {DEEP_THRESHOLD} · {len(deep_entries)} entries",
        deep_entries,
    )

    if not deep_only:
        doc.add_page_break()
        render_section(
            doc,
            "Shallow Discoveries",
            f"Desert depth {SHALLOW_THRESHOLD}–{DEEP_THRESHOLD} · {len(shallow_entries)} entries",
            shallow_entries,
        )

        doc.add_page_break()
        render_section(
            doc,
            "Notes and Manual Entries",
            f"{len(other_entries)} entries",
            other_entries,
        )

        doc.add_page_break()
        render_appendix_a(doc, entries)

        doc.add_page_break()
        render_appendix_b(doc, entries)

    doc.save(output_path)
    total_written = len(deep_entries) if deep_only else len(deep_entries) + len(shallow_entries) + len(other_entries)
    return total_written


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate FIELD_JOURNAL.docx from journal.json.",
    )
    parser.add_argument("--deep-only", action="store_true",
                        help="Generate only FIELD_JOURNAL_DEEP_ONLY.docx (Section 1 + Appendix A)")
    parser.add_argument("--output-dir", default=None,
                        help="Output directory (default: docs/ inside project root)")
    args = parser.parse_args()

    project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    journal_path = os.path.join(project_root, JOURNAL_PATH)

    # Honour --output-dir, keeping output within the project for safety
    global OUTPUT_DIR
    if args.output_dir:
        cand = os.path.abspath(args.output_dir)
        if not cand.startswith(project_root):
            print(f"ERROR: --output-dir must be inside the project directory.")
            sys.exit(1)
        OUTPUT_DIR = os.path.relpath(cand, project_root)

    if not os.path.exists(journal_path):
        print(f"ERROR: journal not found at {journal_path}")
        sys.exit(1)

    with open(journal_path, encoding="utf-8") as f:
        raw = json.load(f)

    # Normalise: accept both a list and {"entries": [...]}
    entries = raw if isinstance(raw, list) else raw.get("entries", [])

    # Validate and filter
    valid, skipped = [], []
    for e in entries:
        if not e.get("id"):
            skipped.append(("no-id", e))
            continue
        dv = e.get("desert_value")
        if dv is not None and not isinstance(dv, (int, float)):
            skipped.append((e["id"], e))
            continue
        valid.append(e)

    if skipped:
        for sid, _ in skipped:
            print(f"  SKIPPED entry: {sid}")

    # Counts for summary
    probe   = [e for e in valid if e.get("type") == "probe_discovery" and isinstance(e.get("desert_value"), float)]
    deep    = [e for e in probe if e["desert_value"] >= DEEP_THRESHOLD]
    shallow = [e for e in probe if SHALLOW_THRESHOLD <= e["desert_value"] < DEEP_THRESHOLD]
    other_set = set(id(e) for e in probe)
    other   = [e for e in valid if id(e) not in other_set or e.get("desert_value", 0) < SHALLOW_THRESHOLD]

    out_full      = safe_output_path("FIELD_JOURNAL.docx")
    out_deep_only = safe_output_path("FIELD_JOURNAL_DEEP_ONLY.docx")
    os.makedirs(os.path.dirname(out_full), exist_ok=True)

    if args.deep_only:
        print("Building FIELD_JOURNAL_DEEP_ONLY.docx …")
        n_deep = build_journal_doc(valid, out_deep_only, deep_only=True)
        size_deep = os.path.getsize(out_deep_only) // 1024
        approx_pages_deep = max(1, n_deep // 3)
        print(f"\nFIELD_JOURNAL_DEEP_ONLY.docx: {n_deep} entries, approx {approx_pages_deep} pages  ({size_deep} KB)")
        return

    print("Building FIELD_JOURNAL.docx …")
    n_full = build_journal_doc(valid, out_full, deep_only=False)

    print("Building FIELD_JOURNAL_DEEP_ONLY.docx …")
    n_deep = build_journal_doc(valid, out_deep_only, deep_only=True)

    size_full = os.path.getsize(out_full) // 1024
    size_deep = os.path.getsize(out_deep_only) // 1024
    approx_pages_full = max(1, n_full // 3)
    approx_pages_deep = max(1, n_deep // 3)

    print()
    print(f"  Deep entries:         {len(deep)}")
    print(f"  Shallow entries:      {len(shallow)}")
    print(f"  Manual/other entries: {len(other)}")
    print()
    print(f"FIELD_JOURNAL.docx:           {n_full} entries, approx {approx_pages_full} pages  ({size_full} KB)")
    print(f"FIELD_JOURNAL_DEEP_ONLY.docx: {n_deep} entries, approx {approx_pages_deep} pages  ({size_deep} KB)")


if __name__ == "__main__":
    main()
