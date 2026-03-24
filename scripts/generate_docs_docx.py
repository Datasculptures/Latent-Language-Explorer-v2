"""
generate_docs_docx.py
Converts READING_THE_TERRAIN_V2.md and SCRIPT_REFERENCE.md to .docx.

Author header, datasculptures.com footer, page numbers.
Run from the project root.
"""

import re
import sys
import os

sys.path.insert(0, "C:/Users/SeanP/OneDrive - mapleclose.ca/Sean/AI/LLEv2/Lib/site-packages")

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AUTHOR = "Sean Patrick Morris · Claude (Anthropic)"
SITE   = "datasculptures.com"
DATE   = "March 2026"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def add_page_numbers(section):
    """Add 'Page N of M' centre footer with page numbers."""
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    run = para.add_run(SITE + "   ·   Page ")
    run.font.size = Pt(9)

    # PAGE field
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)

    run2 = para.add_run(" of ")
    run2.font.size = Pt(9)

    # NUMPAGES field
    run3 = para.add_run()
    run3.font.size = Pt(9)
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "begin")
    run3._r.append(fld3)
    instr2 = OxmlElement("w:instrText")
    instr2.text = "NUMPAGES"
    run3._r.append(instr2)
    fld4 = OxmlElement("w:fldChar")
    fld4.set(qn("w:fldCharType"), "end")
    run3._r.append(fld4)


def set_header(section, title):
    """Add author / title header."""
    header = section.header
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run(f"{AUTHOR}   ·   {title}   ·   {DATE}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def set_margins(section):
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)


# ---------------------------------------------------------------------------
# Paragraph style helpers
# ---------------------------------------------------------------------------

def apply_normal(para):
    para.style = "Normal"
    for run in para.runs:
        run.font.size = Pt(11)


def para_spacing(para, before=0, after=6):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"),  str(after))
    pPr.append(spacing)


# ---------------------------------------------------------------------------
# Inline markdown: **bold**, `code`
# ---------------------------------------------------------------------------

def add_inline_md(para, text):
    """Render **bold** and `code` within a paragraph."""
    # Split on bold or code spans
    tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for tok in tokens:
        if tok.startswith("**") and tok.endswith("**"):
            run = para.add_run(tok[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        elif tok.startswith("`") and tok.endswith("`"):
            run = para.add_run(tok[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x20, 0x60, 0x80)
        else:
            if tok:
                run = para.add_run(tok)
                run.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Table rendering
# ---------------------------------------------------------------------------

def render_table(doc, rows):
    """rows: list of lists of strings (first row = header)."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            add_inline_md(para, cell_text)
            if i == 0:
                for run in para.runs:
                    run.bold = True
            para.runs and None  # noqa: suppress unused expr warning
        # shade header row lightly
        if i == 0:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "DDEEFF")
                tcPr.append(shd)
    doc.add_paragraph()  # breathing room after table


# ---------------------------------------------------------------------------
# Main markdown → docx converter
# ---------------------------------------------------------------------------

def md_to_docx(md_path, docx_path, doc_title):
    doc = Document()
    section = doc.sections[0]
    set_margins(section)
    set_header(section, doc_title)
    add_page_numbers(section)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if table_rows:
            render_table(doc, table_rows)
        in_table = False
        table_rows = []

    def flush_code():
        nonlocal in_code_block, code_lines
        if code_lines:
            text = "\n".join(code_lines)
            para = doc.add_paragraph(style="Normal")
            run = para.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x20, 0x50, 0x20)
            para_spacing(para, before=4, after=4)
        in_code_block = False
        code_lines = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # ---- code block ----
        if line.startswith("```"):
            if in_code_block:
                flush_code()
            else:
                if in_table:
                    flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ---- table ----
        if line.startswith("|"):
            # separator row?
            if re.match(r"^\|[-| :]+\|$", line):
                i += 1
                continue
            cells = [c.strip() for c in line.strip("|").split("|")]
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ---- horizontal rule ----
        if re.match(r"^---+$", line.strip()):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
            para_spacing(para, before=6, after=6)
            i += 1
            continue

        # ---- headings ----
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}.get(level, "Heading 3")
            para = doc.add_paragraph(style=style)
            # strip italic wrapper common in H1 subtitles
            text = re.sub(r"\*([^*]+)\*", r"\1", text)
            add_inline_md(para, text)
            i += 1
            continue

        # ---- bullet list ----
        m = re.match(r"^(\s*)[*\-]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text   = m.group(2)
            para = doc.add_paragraph(style="List Bullet")
            add_inline_md(para, text)
            para_spacing(para, before=0, after=2)
            i += 1
            continue

        # ---- numbered list ----
        m = re.match(r"^\d+\.\s+(.*)", line)
        if m:
            para = doc.add_paragraph(style="List Number")
            add_inline_md(para, m.group(1))
            para_spacing(para, before=0, after=2)
            i += 1
            continue

        # ---- blank line ----
        if line.strip() == "":
            i += 1
            continue

        # ---- italic block (e.g. *datasculptures.com*) ----
        if re.match(r"^\*[^*]", line) and line.endswith("*"):
            text = line.strip("*")
            para = doc.add_paragraph(style="Normal")
            run  = para.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            para_spacing(para, before=4, after=4)
            i += 1
            continue

        # ---- normal paragraph ----
        para = doc.add_paragraph(style="Normal")
        add_inline_md(para, line)
        para_spacing(para, before=0, after=6)
        i += 1

    if in_code_block:
        flush_code()
    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"  Saved: {docx_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

DOCS = [
    (
        "docs/READING_THE_TERRAIN_V2.md",
        "docs/READING_THE_TERRAIN_V2.docx",
        "Reading the Terrain — V2",
    ),
    (
        "docs/SCRIPT_REFERENCE.md",
        "docs/SCRIPT_REFERENCE.docx",
        "Script Reference",
    ),
]

if __name__ == "__main__":
    for md, docx, title in DOCS:
        if not os.path.exists(md):
            print(f"  MISSING: {md}")
            continue
        print(f"Converting {md} …")
        md_to_docx(md, docx, title)

    print("Done.")
